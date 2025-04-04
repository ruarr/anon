import re
import docx
import logging
import os
from pathlib import Path
import phonenumbers
from email_validator import validate_email, EmailNotValidError
import spacy
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Configuración de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    filename='anonimizacion.log'
)

class DocumentAnonymizer:
    def __init__(self):
        self.nlp = spacy.load("es_core_news_lg")
        self.patron_dni = r'\b(?:\d{2}\.\d{3}\.\d{3}|\d{8})\b'
        self.patron_direcciones = r'\b(Calle|Av\.|Avenida|Ruta|Pasaje)\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)?\s+\d+(?:\s*[A-Za-z])?\b'
        self.patron_telefono = r'\+?\d{1,3}[-.\s]?\d{3,4}[-.\s]?\d{4}\b'
        self.patron_email = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

    def anonimizar_texto(self, texto):
        """Anonimiza texto usando IA y regex"""
        try:
            doc = self.nlp(texto)
            for ent in doc.ents:
                if ent.label_ in ["PER", "LOC", "ORG"]:
                    texto = texto.replace(ent.text, f"{ent.label_}_ANONIMIZADO")
            
            texto = re.sub(self.patron_dni, 'DNI_ANONIMIZADO', texto)
            texto = re.sub(self.patron_direcciones, 'DIRECCION_ANONIMIZADA', texto)
            
            def reemplazar_telefono(match):
                try:
                    numero = phonenumbers.parse(match.group(0))
                    if phonenumbers.is_valid_number(numero):
                        return 'TELEFONO_ANONIMIZADO'
                    return match.group(0)
                except phonenumbers.NumberParseException:
                    return match.group(0)
            texto = re.sub(self.patron_telefono, reemplazar_telefono, texto)
            
            def reemplazar_email(match):
                try:
                    validate_email(match.group(0))
                    return 'EMAIL_ANONIMIZADO'
                except EmailNotValidError:
                    return match.group(0)
            texto = re.sub(self.patron_email, reemplazar_email, texto)
            
            return texto
        except Exception as e:
            logging.error(f"Error al anonimizar texto: {str(e)}")
            raise

    def anonimizar_docx(self, archivo_entrada):
        """Lee y anonimiza un archivo DOCX, retorna el documento procesado"""
        try:
            if not os.path.exists(archivo_entrada):
                raise FileNotFoundError(f"El archivo {archivo_entrada} no existe")
            
            doc = Document(archivo_entrada)
            # Procesar párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    anonimizado = self.anonimizar_texto(para.text)
                    para.text = anonimizado  # Reemplazar texto manteniendo estilo
            
            # Procesar tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            anonimizado = self.anonimizar_texto(cell.text)
                            cell.text = anonimizado  # Reemplazar texto manteniendo formato
            
            return doc  # Retornar el documento completo, no solo texto
        except Exception as e:
            logging.error(f"Error al procesar DOCX {archivo_entrada}: {str(e)}")
            raise

    def save_as_docx(self, doc, archivo_salida):
        """Guarda el documento anonimizado como DOCX preservando formato"""
        Path(archivo_salida).parent.mkdir(parents=True, exist_ok=True)
        doc.save(archivo_salida)
        logging.info(f"Archivo DOCX guardado: {archivo_salida}")

    def save_as_txt(self, texto, archivo_salida):
        """Guarda el texto anonimizado como TXT"""
        Path(archivo_salida).parent.mkdir(parents=True, exist_ok=True)
        with open(archivo_salida, 'w', encoding='utf-8') as f:
            f.write(texto)
        logging.info(f"Archivo TXT guardado: {archivo_salida}")

    def save_as_pdf(self, texto, archivo_salida):
        """Guarda el texto anonimizado como PDF usando ReportLab"""
        Path(archivo_salida).parent.mkdir(parents=True, exist_ok=True)
        c = canvas.Canvas(archivo_salida, pagesize=letter)
        width, height = letter
        c.setFont("Helvetica", 12)
        y_position = height - 40
        line_height = 14
        lines = texto.split('\n')
        for line in lines:
            if y_position < 40:
                c.showPage()
                c.setFont("Helvetica", 12)
                y_position = height - 40
            c.drawString(40, y_position, line)
            y_position -= line_height
        c.save()
        logging.info(f"Archivo PDF guardado: {archivo_salida}")

class AnonymizerGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Poder Judicial de Jujuy - INLAB -")
        self.root.geometry("600x400")
        self.anonymizer = DocumentAnonymizer()

        tk.Label(root, text="Archivo de entrada (.docx):").grid(row=0, column=0, padx=10, pady=10)
        self.input_path = tk.Entry(root, width=50)
        self.input_path.grid(row=0, column=1, padx=10, pady=10)
        tk.Button(root, text="Seleccionar", command=self.select_input_file).grid(row=0, column=2, padx=10, pady=10)

        tk.Label(root, text="Archivo de salida:").grid(row=1, column=0, padx=10, pady=10)
        self.output_path = tk.Entry(root, width=50)
        self.output_path.grid(row=1, column=1, padx=10, pady=10)
        tk.Button(root, text="Seleccionar", command=self.select_output_file).grid(row=1, column=2, padx=10, pady=10)

        tk.Label(root, text="Formato de salida:").grid(row=2, column=0, padx=10, pady=10)
        self.output_format = tk.StringVar(value="DOCX")
        self.output_format.trace("w", self.update_output_suggestion)
        tk.Radiobutton(root, text="DOCX", variable=self.output_format, value="DOCX").grid(row=2, column=1, padx=5, pady=10, sticky="w")
        tk.Radiobutton(root, text="TXT", variable=self.output_format, value="TXT").grid(row=2, column=1, padx=5, pady=10)
        tk.Radiobutton(root, text="PDF", variable=self.output_format, value="PDF").grid(row=2, column=1, padx=5, pady=10, sticky="e")

        tk.Button(root, text="Anonimizar", command=self.anonimizar).grid(row=3, column=1, pady=20)

    def select_input_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
        if file_path:
            self.input_path.delete(0, tk.END)
            self.input_path.insert(0, file_path)
            self.update_output_suggestion()

    def select_output_file(self):
        format_extensions = {
            "DOCX": ".docx",
            "TXT": ".txt",
            "PDF": ".pdf"
        }
        selected_format = self.output_format.get()
        file_path = filedialog.asksaveasfilename(
            defaultextension=format_extensions[selected_format],
            filetypes=[(f"Archivos {selected_format}", f"*{format_extensions[selected_format]}")]
        )
        if file_path:
            self.output_path.delete(0, tk.END)
            self.output_path.insert(0, file_path)

    def update_output_suggestion(self, *args):
        input_file = self.input_path.get()
        if input_file:
            base_name = os.path.splitext(input_file)[0]
            format_extensions = {
                "DOCX": "_anonimizado.docx",
                "TXT": "_anonimizado.txt",
                "PDF": "_anonimizado.pdf"
            }
            selected_format = self.output_format.get()
            output_path = base_name + format_extensions[selected_format]
            self.output_path.delete(0, tk.END)
            self.output_path.insert(0, output_path)

    def anonimizar(self):
        input_file = self.input_path.get()
        output_file = self.output_path.get()
        output_format = self.output_format.get()

        if not input_file or not output_file:
            messagebox.showerror("Error", "Por favor, seleccione los archivos de entrada y salida.")
            return

        try:
            if output_format == "DOCX":
                doc_anonimizado = self.anonymizer.anonimizar_docx(input_file)
                self.anonymizer.save_as_docx(doc_anonimizado, output_file)
            else:
                # Para TXT y PDF, extraemos el texto plano del documento anonimizado
                doc = self.anonymizer.anonimizar_docx(input_file)
                texto_anonimizado = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                if output_format == "TXT":
                    self.anonymizer.save_as_txt(texto_anonimizado, output_file)
                elif output_format == "PDF":
                    self.anonymizer.save_as_pdf(texto_anonimizado, output_file)
            
            messagebox.showinfo("Éxito", f"El archivo ha sido anonimizado y guardado como {output_format}.")
        except Exception as e:
            messagebox.showerror("Error", f"Error al anonimizar el archivo: {str(e)}")

def main():
    root = tk.Tk()
    app = AnonymizerGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()